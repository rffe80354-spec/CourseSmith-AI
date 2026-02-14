"""
DOCX Engine Module - Professional Word Document creation for CourseSmith AI.
Uses python-docx for generating editable Microsoft Word documents.
SECURITY: Requires valid session token to function (anti-tamper protection).
"""

import os
import re
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from session_manager import is_active, get_tier, SecurityError


class DocxBuilder:
    """
    Professional DOCX document builder with styling and formatting.
    
    Features:
    - Cover page with title and subtitle
    - Chapter headers with proper formatting
    - Markdown parsing (# and ##, bullet points)
    - Clean typography
    - Table of Contents support
    """

    def __init__(self, filename, tier=None):
        """
        Initialize the DOCX builder with document settings.

        Args:
            filename: The output DOCX filename.
            tier: License tier (if None, fetched from session).
        """
        # Security check
        if not is_active():
            raise SecurityError("Unauthorized: No valid session. Please activate your license.")
        
        self.filename = filename
        self.tier = tier if tier is not None else (get_tier() or 'trial')
        self.doc = Document()
        self._setup_styles()
        self.title = ""
        self.subtitle = ""

    def _setup_styles(self):
        """Setup custom styles for the document."""
        styles = self.doc.styles
        
        # Title style
        if 'CourseTitle' not in [s.name for s in styles]:
            title_style = styles.add_style('CourseTitle', WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.name = 'Arial'
            title_style.font.size = Pt(36)
            title_style.font.bold = True
            title_style.font.color.rgb = RGBColor(26, 26, 46)
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_style.paragraph_format.space_after = Pt(24)
        
        # Subtitle style
        if 'CourseSubtitle' not in [s.name for s in styles]:
            sub_style = styles.add_style('CourseSubtitle', WD_STYLE_TYPE.PARAGRAPH)
            sub_style.font.name = 'Arial'
            sub_style.font.size = Pt(18)
            sub_style.font.color.rgb = RGBColor(74, 74, 106)
            sub_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sub_style.paragraph_format.space_after = Pt(48)
        
        # Chapter header style
        if 'ChapterHeader' not in [s.name for s in styles]:
            ch_style = styles.add_style('ChapterHeader', WD_STYLE_TYPE.PARAGRAPH)
            ch_style.font.name = 'Arial'
            ch_style.font.size = Pt(24)
            ch_style.font.bold = True
            ch_style.font.color.rgb = RGBColor(26, 26, 46)
            ch_style.paragraph_format.space_before = Pt(24)
            ch_style.paragraph_format.space_after = Pt(12)
        
        # Section header style
        if 'SectionHeader' not in [s.name for s in styles]:
            sec_style = styles.add_style('SectionHeader', WD_STYLE_TYPE.PARAGRAPH)
            sec_style.font.name = 'Arial'
            sec_style.font.size = Pt(16)
            sec_style.font.bold = True
            sec_style.font.color.rgb = RGBColor(54, 54, 86)
            sec_style.paragraph_format.space_before = Pt(18)
            sec_style.paragraph_format.space_after = Pt(8)
        
        # Body text style
        if 'CourseBody' not in [s.name for s in styles]:
            body_style = styles.add_style('CourseBody', WD_STYLE_TYPE.PARAGRAPH)
            body_style.font.name = 'Arial'
            body_style.font.size = Pt(11)
            body_style.font.color.rgb = RGBColor(33, 33, 33)
            body_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            body_style.paragraph_format.space_after = Pt(10)

    def set_title(self, title, subtitle=""):
        """
        Set the document title and subtitle.
        
        Args:
            title: Main title of the course.
            subtitle: Subtitle or description.
        """
        self.title = title
        self.subtitle = subtitle

    def add_cover_page(self, title=None, subtitle=None):
        """
        Add a cover page to the document.
        
        Args:
            title: Main title (uses set_title value if None).
            subtitle: Subtitle (uses set_title value if None).
        """
        if title:
            self.title = title
        if subtitle:
            self.subtitle = subtitle
        
        # Add spacing at top
        for _ in range(8):
            self.doc.add_paragraph()
        
        # Add title
        title_para = self.doc.add_paragraph(self.title, style='CourseTitle')
        
        # Add subtitle
        if self.subtitle:
            self.doc.add_paragraph(self.subtitle, style='CourseSubtitle')
        
        # Add page break after cover
        self.doc.add_page_break()

    def add_table_of_contents(self, chapters):
        """
        Add a table of contents.
        
        Args:
            chapters: List of chapter titles.
        """
        # TOC Title
        toc_title = self.doc.add_paragraph("Table of Contents")
        toc_title.style = 'ChapterHeader'
        
        # Add chapter entries
        for i, chapter in enumerate(chapters, 1):
            toc_entry = self.doc.add_paragraph()
            toc_entry.add_run(f"Chapter {i}: ").bold = True
            toc_entry.add_run(chapter)
            toc_entry.paragraph_format.space_after = Pt(6)
        
        self.doc.add_page_break()

    def _parse_markdown_line(self, line):
        """
        Parse a line for markdown formatting.
        
        Args:
            line: Text line to parse.
            
        Returns:
            Tuple of (text, style, is_bullet)
        """
        line = line.strip()
        
        # Check for headers
        if line.startswith('# '):
            return line[2:], 'ChapterHeader', False
        elif line.startswith('## '):
            return line[3:], 'SectionHeader', False
        
        # Check for bullet points
        if line.startswith('- ') or line.startswith('* '):
            return line[2:], 'CourseBody', True
        if line.startswith('• '):
            return line[1:].lstrip(), 'CourseBody', True
        
        return line, 'CourseBody', False

    def add_chapter(self, title, content, chapter_num=None):
        """
        Add a chapter to the document.
        
        Args:
            title: Chapter title.
            content: Chapter content (markdown supported).
            chapter_num: Optional chapter number.
        """
        # Chapter header
        if chapter_num:
            header_text = f"Chapter {chapter_num}: {title}"
        else:
            header_text = title
        
        chapter_header = self.doc.add_paragraph(header_text, style='ChapterHeader')
        
        # Parse and add content
        lines = content.split('\n')
        
        for line in lines:
            if not line.strip():
                continue
            
            text, style, is_bullet = self._parse_markdown_line(line)
            
            if not text:
                continue
            
            if is_bullet:
                # Add bullet point
                para = self.doc.add_paragraph(text, style='List Bullet')
            else:
                # Add regular paragraph
                para = self.doc.add_paragraph(text, style=style)
        
        # Add spacing after chapter
        self.doc.add_paragraph()

    def add_paragraph(self, text, style='CourseBody'):
        """
        Add a paragraph with the specified style.
        
        Args:
            text: Paragraph text.
            style: Style name to apply.
        """
        self.doc.add_paragraph(text, style=style)

    def build(self):
        """
        Build and save the document.
        
        Returns:
            str: Path to the saved document.
        """
        # Security check before saving
        if not is_active():
            raise SecurityError("Unauthorized: Session expired. Please re-activate your license.")
        
        self.doc.save(self.filename)
        return self.filename


def create_course_docx(filename, title, subtitle, chapters, contents):
    """
    Create a complete course DOCX document.
    
    Args:
        filename: Output filename.
        title: Course title.
        subtitle: Course subtitle.
        chapters: List of chapter titles.
        contents: List of chapter contents (same order as chapters).
        
    Returns:
        str: Path to the saved document.
    """
    builder = DocxBuilder(filename)
    builder.set_title(title, subtitle)
    builder.add_cover_page()
    builder.add_table_of_contents(chapters)
    
    for i, (chapter_title, chapter_content) in enumerate(zip(chapters, contents), 1):
        builder.add_chapter(chapter_title, chapter_content, chapter_num=i)
    
    return builder.build()
