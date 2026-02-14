"""
DOCX Exporter Module - Export CourseSmith content to Microsoft Word format.
Uses python-docx library for professional DOCX generation.
"""

import os
import re
from datetime import datetime
from typing import Optional

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from export_base import ExporterBase, ExportError, ExportManager


class DOCXExporter(ExporterBase):
    """
    Exporter for Microsoft Word DOCX format.
    
    Creates professional Word documents with:
    - Cover page with title and subtitle
    - Table of contents
    - Formatted chapter content
    - Headers and footers
    """
    
    file_extension = "docx"
    format_name = "Microsoft Word"
    
    def __init__(self, project, output_path: str = None):
        """
        Initialize the DOCX exporter.
        
        Args:
            project: CourseProject object with content to export.
            output_path: Optional output file path.
        """
        super().__init__(project, output_path)
        
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx is required for DOCX export. "
                "Install it with: pip install python-docx"
            )
        
        self.doc = None
    
    def export(self) -> str:
        """
        Export the project to DOCX format.
        
        Returns:
            str: Path to the exported DOCX file.
            
        Raises:
            ExportError: If export fails.
        """
        if not self.validate_project():
            raise ExportError(f"Invalid project: {', '.join(self.errors)}")
        
        try:
            # Create document
            self.doc = Document()
            
            # Setup custom styles
            self._setup_styles()
            
            # Add cover page
            self._add_cover_page()
            
            # Add table of contents placeholder
            self._add_toc()
            
            # Add chapters
            self._add_chapters()
            
            # Generate output path if needed
            output_path = self.generate_output_path()
            
            # Save document
            self.doc.save(output_path)
            
            return output_path
            
        except Exception as e:
            raise ExportError(f"Failed to export DOCX: {str(e)}")
    
    def _setup_styles(self):
        """Setup custom document styles."""
        styles = self.doc.styles
        
        # Title style
        try:
            title_style = styles.add_style('CourseTitle', WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.size = Pt(36)
            title_style.font.bold = True
            title_style.font.color.rgb = RGBColor(26, 26, 46)
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_style.paragraph_format.space_after = Pt(24)
        except ValueError:
            pass  # Style already exists
        
        # Subtitle style
        try:
            subtitle_style = styles.add_style('CourseSubtitle', WD_STYLE_TYPE.PARAGRAPH)
            subtitle_style.font.size = Pt(18)
            subtitle_style.font.color.rgb = RGBColor(74, 74, 106)
            subtitle_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle_style.paragraph_format.space_after = Pt(12)
        except ValueError:
            pass
        
        # Chapter heading style
        try:
            chapter_style = styles.add_style('ChapterHeading', WD_STYLE_TYPE.PARAGRAPH)
            chapter_style.font.size = Pt(24)
            chapter_style.font.bold = True
            chapter_style.font.color.rgb = RGBColor(26, 26, 46)
            chapter_style.paragraph_format.space_before = Pt(24)
            chapter_style.paragraph_format.space_after = Pt(12)
        except ValueError:
            pass
        
        # Section heading style
        try:
            section_style = styles.add_style('SectionHeading', WD_STYLE_TYPE.PARAGRAPH)
            section_style.font.size = Pt(16)
            section_style.font.bold = True
            section_style.font.color.rgb = RGBColor(42, 42, 78)
            section_style.paragraph_format.space_before = Pt(18)
            section_style.paragraph_format.space_after = Pt(8)
        except ValueError:
            pass
    
    def _add_cover_page(self):
        """Add a cover page with title and metadata."""
        # Add some spacing
        for _ in range(3):
            self.doc.add_paragraph()
        
        # Add cover image if available
        if self.project.cover_image_path and os.path.exists(self.project.cover_image_path):
            try:
                self.doc.add_picture(self.project.cover_image_path, width=Inches(5))
                last_para = self.doc.paragraphs[-1]
                last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                self.doc.add_paragraph()
            except Exception:
                pass  # Skip image if it fails
        
        # Title
        try:
            title = self.doc.add_paragraph(self.project.topic, style='CourseTitle')
        except KeyError:
            title = self.doc.add_paragraph(self.project.topic)
            title.runs[0].font.size = Pt(36)
            title.runs[0].font.bold = True
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtitle
        subtitle_text = f"A Comprehensive Guide for {self.project.audience}"
        try:
            subtitle = self.doc.add_paragraph(subtitle_text, style='CourseSubtitle')
        except KeyError:
            subtitle = self.doc.add_paragraph(subtitle_text)
            subtitle.runs[0].font.size = Pt(18)
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add spacing
        for _ in range(4):
            self.doc.add_paragraph()
        
        # Author/Company info
        branding = self.project.branding
        if branding.get('author_name'):
            author = self.doc.add_paragraph(f"By {branding['author_name']}")
            author.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if branding.get('company_name'):
            company = self.doc.add_paragraph(branding['company_name'])
            company.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Generation info
        self.doc.add_paragraph()
        gen_info = self.doc.add_paragraph("Generated by CourseSmith AI")
        gen_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        gen_info.runs[0].font.size = Pt(10)
        gen_info.runs[0].font.color.rgb = RGBColor(128, 128, 128)
        
        # Page break
        self.doc.add_page_break()
    
    def _add_toc(self):
        """Add table of contents."""
        toc_heading = self.doc.add_paragraph("Table of Contents")
        toc_heading.runs[0].font.size = Pt(24)
        toc_heading.runs[0].font.bold = True
        
        self.doc.add_paragraph()
        
        # Add chapter links
        for i, title in enumerate(self.project.outline, 1):
            toc_entry = self.doc.add_paragraph()
            toc_entry.add_run(f"Chapter {i}: {title}")
            toc_entry.paragraph_format.left_indent = Inches(0.5)
        
        self.doc.add_page_break()
    
    def _add_chapters(self):
        """Add chapter content."""
        for i, chapter_title in enumerate(self.project.outline, 1):
            # Chapter heading
            try:
                heading = self.doc.add_paragraph(
                    f"Chapter {i}: {chapter_title}", 
                    style='ChapterHeading'
                )
            except KeyError:
                heading = self.doc.add_paragraph(f"Chapter {i}: {chapter_title}")
                heading.runs[0].font.size = Pt(24)
                heading.runs[0].font.bold = True
            
            # Chapter content
            content = self.project.chapters_content.get(chapter_title, "")
            self._add_formatted_content(content)
            
            # Page break after each chapter (except last)
            if i < len(self.project.outline):
                self.doc.add_page_break()
    
    def _add_formatted_content(self, content: str):
        """
        Add formatted content with markdown parsing.
        
        Args:
            content: Markdown-formatted content string.
        """
        lines = content.split('\n')
        current_para_text = []
        
        for line in lines:
            stripped = line.strip()
            
            if stripped.startswith('## '):
                # Flush current paragraph
                if current_para_text:
                    self._add_paragraph(' '.join(current_para_text))
                    current_para_text = []
                
                # Add section header
                header_text = stripped[3:].strip()
                try:
                    self.doc.add_paragraph(header_text, style='SectionHeading')
                except KeyError:
                    p = self.doc.add_paragraph(header_text)
                    p.runs[0].font.size = Pt(16)
                    p.runs[0].font.bold = True
                    
            elif stripped.startswith('# '):
                # Flush current paragraph
                if current_para_text:
                    self._add_paragraph(' '.join(current_para_text))
                    current_para_text = []
                
                # Add main header
                header_text = stripped[1:].strip()
                p = self.doc.add_paragraph(header_text)
                p.runs[0].font.size = Pt(20)
                p.runs[0].font.bold = True
                    
            elif stripped.startswith('* ') or stripped.startswith('- '):
                # Flush current paragraph
                if current_para_text:
                    self._add_paragraph(' '.join(current_para_text))
                    current_para_text = []
                
                # Add bullet point
                bullet_text = stripped[2:].strip()
                self._add_bullet_point(bullet_text)
                
            elif stripped == '':
                # Empty line - flush paragraph
                if current_para_text:
                    self._add_paragraph(' '.join(current_para_text))
                    current_para_text = []
            else:
                # Regular text
                current_para_text.append(stripped)
        
        # Flush remaining paragraph
        if current_para_text:
            self._add_paragraph(' '.join(current_para_text))
    
    def _add_paragraph(self, text: str):
        """Add a formatted paragraph with bold/italic support."""
        para = self.doc.add_paragraph()
        
        # Process markdown bold and italic
        # Split on **text** patterns
        parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', text)
        
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                # Bold text
                run = para.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*'):
                # Italic text
                run = para.add_run(part[1:-1])
                run.italic = True
            else:
                para.add_run(part)
    
    def _add_bullet_point(self, text: str):
        """Add a bullet point item."""
        para = self.doc.add_paragraph(text, style='List Bullet')


# Register the exporter
ExportManager.register_exporter('docx', DOCXExporter)
